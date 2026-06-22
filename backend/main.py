from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import uvicorn
import tempfile, os, re, difflib
from typing import Optional

# Document parsers
import fitz  # PyMuPDF
from docx import Document

app = FastAPI(title="Publishing QA Validation API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Helpers ────────────────────────────────────────────────────────────────

def extract_docx(path: str) -> dict:
    doc = Document(path)
    paragraphs = []
    headings   = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        entry = {"index": i, "text": text, "style": style}
        paragraphs.append(entry)
        if "Heading" in style:
            headings.append(entry)
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append({"table_index": ti + 1, "rows": rows})
    full_text = "\n".join(p["text"] for p in paragraphs)
    return {"paragraphs": paragraphs, "headings": headings,
            "tables": tables, "full_text": full_text}


def extract_pdf(path: str) -> dict:
    doc   = fitz.open(path)
    pages = []
    for i, page in enumerate(doc):
        text   = page.get_text("text").strip()
        blocks = page.get_text("blocks")
        pages.append({
            "page_num" : i + 1,
            "text"     : text,
            "blocks"   : [b[4] for b in blocks if b[4].strip()],
        })
    full_text = "\n".join(p["text"] for p in pages)
    return {"pages": pages, "full_text": full_text, "total_pages": len(pages)}


# ─── Individual check functions ─────────────────────────────────────────────

def check_page_sequence(pdf_data):
    errors = []
    pages  = pdf_data["pages"]
    for page in pages:
        text    = page["text"]
        numbers = re.findall(r'(?<!\d)(\d{1,4})(?!\d)', text)
        nums    = [int(n) for n in numbers if 1 <= int(n) <= 9999]
        if nums:
            found  = nums[0]
            expect = page["page_num"]
            if abs(found - expect) > 3:
                errors.append({
                    "check"   : "Page Number Sequence & Folio",
                    "page"    : str(page["page_num"]),
                    "location": f"Folio shows {found}, expected ~{expect}",
                })
    return errors


def check_running_heads(pdf_data):
    errors = []
    for page in pdf_data["pages"]:
        lines = [l.strip() for l in page["text"].split("\n") if l.strip()]
        if not lines:
            continue
        head = lines[0]
        if len(head) > 120:
            errors.append({
                "check"   : "Running Head Style & Position",
                "page"    : str(page["page_num"]),
                "location": f"Running head may be truncated or mis-positioned: '{head[:60]}…'",
            })
    return errors


def check_slug_line(pdf_data):
    errors = []
    slug_pattern = re.compile(r'\b(Ch|Chapter|Sec|Section|Fig|Table)\b', re.I)
    for page in pdf_data["pages"]:
        lines = [l.strip() for l in page["text"].split("\n") if l.strip()]
        if lines and not slug_pattern.search(lines[0]):
            errors.append({
                "check"   : "Slug Line Page Range & File Name",
                "page"    : str(page["page_num"]),
                "location": f"Slug may be missing or malformed: '{lines[0][:60]}'",
            })
    return errors


def check_word_comparison(word_data, pdf_data):
    errors = []
    word_paras = [p["text"] for p in word_data["paragraphs"]]
    pdf_text   = pdf_data["full_text"]
    for i, para in enumerate(word_paras):
        if len(para) < 20:
            continue
        # Check if a significant chunk is missing from PDF
        snippet = para[:60]
        if snippet not in pdf_text:
            ratio = difflib.SequenceMatcher(None, para, pdf_text).ratio()
            if ratio < 0.55:
                errors.append({
                    "check"   : "Word-to-Word Comparison",
                    "page"    : "—",
                    "location": f"Paragraph {i+1}: '{para[:80]}…' not found in PDF",
                })
    return errors[:8]   # cap at 8


def check_typos(word_data, pdf_data):
    errors = []
    common_typos = {
        "teh": "the", "recieve": "receive", "occured": "occurred",
        "seperete": "separate", "definately": "definitely",
        "accomodate": "accommodate", "untill": "until",
        "publsihed": "published", "refernce": "reference",
        "Figuure": "Figure", "Tablel": "Table",
    }
    for src, correction in common_typos.items():
        pattern = re.compile(r'\b' + re.escape(src) + r'\b', re.I)
        for page in pdf_data["pages"]:
            if pattern.search(page["text"]):
                errors.append({
                    "check"   : "Typos",
                    "page"    : str(page["page_num"]),
                    "location": f"Possible typo '{src}' (should be '{correction}')",
                })
    return errors


def check_missing_content(word_data, pdf_data):
    errors = []
    pdf_text = pdf_data["full_text"].lower()
    for heading in word_data["headings"]:
        h = heading["text"]
        if len(h) < 4:
            continue
        if h.lower() not in pdf_text:
            errors.append({
                "check"   : "Missing Content",
                "page"    : "—",
                "location": f"Heading not found in PDF: '{h}'",
            })
    return errors[:6]


def check_content_order(word_data, pdf_data):
    errors = []
    pdf_text = pdf_data["full_text"]
    positions = []
    for heading in word_data["headings"]:
        h   = heading["text"]
        pos = pdf_text.find(h[:40])
        positions.append((h, pos))
    for i in range(1, len(positions)):
        if positions[i][1] != -1 and positions[i-1][1] != -1:
            if positions[i][1] < positions[i-1][1]:
                errors.append({
                    "check"   : "Content Order",
                    "page"    : "—",
                    "location": f"'{positions[i][0][:60]}' appears before '{positions[i-1][0][:60]}' in PDF",
                })
    return errors


def check_headings(word_data, pdf_data):
    errors = []
    pdf_text = pdf_data["full_text"]
    num_re   = re.compile(r'^(\d+[\.\d]*)\s')
    heading_nums = []
    for h in word_data["headings"]:
        m = num_re.match(h["text"])
        if m:
            heading_nums.append(m.group(1))
    for num in heading_nums:
        if num not in pdf_text:
            errors.append({
                "check"   : "Heading Levels & Numbering",
                "page"    : "—",
                "location": f"Heading number '{num}' not found in PDF",
            })
    return errors[:5]


def check_equations(word_data, pdf_data):
    errors = []
    eq_re = re.compile(r'[A-Za-z]\s*=\s*[\w\(\)\+\-\*\/\^]+')
    for page in pdf_data["pages"]:
        matches = eq_re.findall(page["text"])
        for m in matches:
            if "?" in m or "□" in m or "■" in m:
                errors.append({
                    "check"   : "Equations",
                    "page"    : str(page["page_num"]),
                    "location": f"Equation may have rendering issue: '{m[:60]}'",
                })
    return errors


def check_special_chars(pdf_data):
    errors = []
    bad = re.compile(r'[□■▲▼◆●○◎※†‡¶§©®™…–—]')
    for page in pdf_data["pages"]:
        found = bad.findall(page["text"])
        if found:
            unique = list(set(found))
            errors.append({
                "check"   : "Special Characters & Symbols",
                "page"    : str(page["page_num"]),
                "location": f"Possibly garbled symbols: {' '.join(unique[:8])}",
            })
    return errors[:4]


def check_footnotes(word_data, pdf_data):
    errors = []
    fn_re = re.compile(r'\[\d+\]|\(\d+\)|(?<!\d)\d{1,2}(?!\d)\s*$', re.M)
    for page in pdf_data["pages"]:
        refs = fn_re.findall(page["text"])
        for r in refs:
            num = re.search(r'\d+', r)
            if num and int(num.group()) > 200:
                errors.append({
                    "check"   : "Footnote Citation & Placement",
                    "page"    : str(page["page_num"]),
                    "location": f"Footnote reference {r.strip()} seems out of range",
                })
    return errors[:4]


def check_unwanted_chars(pdf_data):
    errors = []
    patterns = {
        r'\$\$\$'          : "$$$",
        r'\bxxx\b'         : "xxx",
        r'\bXXX\b'         : "XXX",
        r'\b000\b'         : "000",
        r'\bTBD\b'         : "TBD",
        r'\bLorem\s+Ipsum\b': "Lorem Ipsum",
        r'\bFPO\b'         : "FPO",
        r'\bPLACEHOLDER\b' : "PLACEHOLDER",
    }
    for pattern, label in patterns.items():
        rx = re.compile(pattern, re.I)
        for page in pdf_data["pages"]:
            if rx.search(page["text"]):
                errors.append({
                    "check"   : "Unwanted Characters",
                    "page"    : str(page["page_num"]),
                    "location": f"Found unwanted placeholder text: '{label}'",
                })
    return errors


def check_citations(pdf_data):
    errors = []
    cit_re = re.compile(r'\(([A-Z][a-z]+(?:,?\s(?:et al\.?|&|and)\s[A-Z][a-z]+)?,?\s*\d{4})\)')
    for page in pdf_data["pages"]:
        citations = cit_re.findall(page["text"])
        for c in citations:
            # check citation is near sentence end
            ctx = page["text"]
            idx = ctx.find(c)
            if idx > 0:
                before = ctx[max(0,idx-5):idx]
                if re.search(r'[A-Z]$', before.strip()):
                    errors.append({
                        "check"   : "Citations & Placement",
                        "page"    : str(page["page_num"]),
                        "location": f"Citation '{c}' may be mid-sentence",
                    })
    return errors[:4]


def check_lists(pdf_data):
    errors = []
    list_re  = re.compile(r'^(\s*[\•\-\*]\s|\s*\d+\.\s)', re.M)
    blank_re = re.compile(r'\n{3,}')
    for page in pdf_data["pages"]:
        if list_re.search(page["text"]) and blank_re.search(page["text"]):
            errors.append({
                "check"   : "List Spacing",
                "page"    : str(page["page_num"]),
                "location": "Excessive blank lines detected near list items",
            })
    return errors[:3]


def check_font_consistency(pdf_data):
    errors = []
    # Heuristic: mixed ALL-CAPS and title-case headings on same page
    allcaps = re.compile(r'\b[A-Z]{4,}\b')
    for page in pdf_data["pages"]:
        caps_count = len(allcaps.findall(page["text"]))
        if caps_count > 10:
            errors.append({
                "check"   : "Font Consistency",
                "page"    : str(page["page_num"]),
                "location": f"High density of ALL-CAPS text ({caps_count} instances) — possible font override",
            })
    return errors[:3]


def check_quotations(word_data, pdf_data):
    errors = []
    quote_re = re.compile(r'["\u201c\u201d]([^""\u201c\u201d]{10,300})["\u201c\u201d]')
    word_quotes = quote_re.findall(word_data["full_text"])
    pdf_text    = pdf_data["full_text"]
    for q in word_quotes[:20]:
        if q[:30] not in pdf_text:
            errors.append({
                "check"   : "Quotations",
                "page"    : "—",
                "location": f"Quote not found in PDF: '\"{q[:60]}\"'",
            })
    return errors[:4]


def check_fpo(pdf_data):
    errors = []
    fpo_re = re.compile(r'\bFPO\b|\bFor\s+Position\s+Only\b|\bPLACEHOLDER\s+IMAGE\b', re.I)
    for page in pdf_data["pages"]:
        if fpo_re.search(page["text"]):
            errors.append({
                "check"   : "FPO / Placeholder Images",
                "page"    : str(page["page_num"]),
                "location": "FPO or placeholder image marker found",
            })
    return errors


def check_tables(word_data, pdf_data):
    errors = []
    pdf_text = pdf_data["full_text"]
    for ti, table in enumerate(word_data["tables"]):
        if not table["rows"]:
            continue
        header_cell = table["rows"][0][0] if table["rows"][0] else ""
        if header_cell and header_cell[:20] not in pdf_text:
            errors.append({
                "check"   : "Table Style Consistency",
                "page"    : "—",
                "location": f"Table {ti+1} header cell '{header_cell[:40]}' not found in PDF",
            })
    return errors[:4]


# ─── Main validation endpoint ───────────────────────────────────────────────

@app.post("/validate")
async def validate(
    word_file: UploadFile = File(...),
    pdf_file : UploadFile = File(...),
    checks   : str        = "",
):
    selected = [c.strip() for c in checks.split(",") if c.strip()] if checks else []

    # Save uploads to temp files
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as wf:
        wf.write(await word_file.read())
        word_path = wf.name

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pf:
        pf.write(await pdf_file.read())
        pdf_path = pf.name

    try:
        word_data = extract_docx(word_path)
        pdf_data  = extract_pdf(pdf_path)
    except Exception as e:
        os.unlink(word_path)
        os.unlink(pdf_path)
        raise HTTPException(status_code=400, detail=f"Failed to parse files: {str(e)}")

    os.unlink(word_path)
    os.unlink(pdf_path)

    # Run all check functions
    all_errors = []

    def run_if_selected(check_name, fn):
        if not selected or check_name in selected:
            try:
                result = fn()
                all_errors.extend(result)
            except Exception:
                pass

    run_if_selected("Page Number Sequence & Folio",  lambda: check_page_sequence(pdf_data))
    run_if_selected("Running Head Style & Position",  lambda: check_running_heads(pdf_data))
    run_if_selected("Slug Line Page Range & File Name", lambda: check_slug_line(pdf_data))
    run_if_selected("Word-to-Word Comparison",        lambda: check_word_comparison(word_data, pdf_data))
    run_if_selected("Typos",                          lambda: check_typos(word_data, pdf_data))
    run_if_selected("Missing Content",                lambda: check_missing_content(word_data, pdf_data))
    run_if_selected("Content Order",                  lambda: check_content_order(word_data, pdf_data))
    run_if_selected("Heading Levels & Numbering",     lambda: check_headings(word_data, pdf_data))
    run_if_selected("Equations",                      lambda: check_equations(word_data, pdf_data))
    run_if_selected("Special Characters & Symbols",   lambda: check_special_chars(pdf_data))
    run_if_selected("Footnote Citation & Placement",  lambda: check_footnotes(word_data, pdf_data))
    run_if_selected("List Spacing",                   lambda: check_lists(pdf_data))
    run_if_selected("Font Consistency",               lambda: check_font_consistency(pdf_data))
    run_if_selected("Quotations",                     lambda: check_quotations(word_data, pdf_data))
    run_if_selected("Citations & Placement",          lambda: check_citations(pdf_data))
    run_if_selected("Unwanted Characters",            lambda: check_unwanted_chars(pdf_data))
    run_if_selected("FPO / Placeholder Images",       lambda: check_fpo(pdf_data))
    run_if_selected("Table Style Consistency",        lambda: check_tables(word_data, pdf_data))

    # Summary
    affected_pages = sorted(set(
        int(e["page"]) for e in all_errors
        if e.get("page", "").isdigit()
    ))

    return {
        "errors"          : all_errors,
        "total_errors"    : len(all_errors),
        "total_pages"     : pdf_data["total_pages"],
        "affected_pages"  : affected_pages,
        "checks_run"      : len(selected) if selected else 18,
    }


@app.get("/health")
def health():
    return {"status": "ok"}


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
